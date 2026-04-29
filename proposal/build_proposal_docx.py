"""Build the SmartTicket project proposal as a .docx file."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(__file__).parent / "SmartTicket_Project_Proposal.docx"


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Calibri"
        if level == 0:
            run.font.size = Pt(26)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x4E, 0x7E)


def _add_paragraph(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
                   align=None, size: int = 11) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               col_widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, "1F3A5F")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths_cm is not None:
        for row in table.rows:
            for c_idx, width in enumerate(col_widths_cm):
                row.cells[c_idx].width = Cm(width)

    doc.add_paragraph()


def build() -> Path:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SmartTicket")
    run.bold = True
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("An AI-Powered Pipeline for Automated\nSupport Ticket Classification and Triage")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x4E, 0x7E)
    run.italic = True

    for _ in range(2):
        doc.add_paragraph()

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run("Project Proposal")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Team table on cover
    team_rows = [
        ["Yousef Yasser",   "221101191"],
        ["Belal Elgendy",   "221101165"],
        ["Ahmed Orabi",     "221101196"],
        ["Mohamed Yousry",  "221101215"],
    ]
    team_table = doc.add_table(rows=1 + len(team_rows), cols=2)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    team_table.style = "Light Grid Accent 1"

    hdr = team_table.rows[0]
    for idx, header in enumerate(["Team Member", "Student ID"]):
        cell = hdr.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        _set_cell_shading(cell, "1F3A5F")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, (name, sid) in enumerate(team_rows, start=1):
        for c_idx, value in enumerate([name, sid]):
            cell = team_table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.size = Pt(11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in team_table.rows:
        row.cells[0].width = Cm(7.5)
        row.cells[1].width = Cm(4.5)

    doc.add_page_break()

    # 1. Abstract
    _add_heading(doc, "1. Abstract", level=1)
    _add_paragraph(
        doc,
        "SmartTicket is an end-to-end machine-learning system that automates the "
        "classification of customer support tickets by department and priority. The "
        "pipeline ingests raw, intentionally noisy ticket data from a relational "
        "database, applies layered preprocessing across text, numeric, and categorical "
        "modalities, and engineers domain-specific features that capture customer "
        "behaviour, ticket linguistics, and operational metadata. Two complementary "
        "feature-selection strategies are then employed: a deterministic greedy-forward "
        "heuristic and a stochastic genetic algorithm. Selected features feed both "
        "voting and stacking ensemble classifiers, and every prediction is accompanied "
        "by word-level explanations produced by LIME and SHAP. The result is a "
        "transparent, reproducible triage system that reduces manual routing effort "
        "while remaining auditable to support agents and supervisors."
    )

    # 2. Problem Statement
    _add_heading(doc, "2. Problem Statement and Motivation", level=1)
    _add_paragraph(
        doc,
        "Customer support teams at mid-to-large organisations receive hundreds — often "
        "thousands — of tickets every day across email, chat, phone, social media, and "
        "web forms. Each incoming ticket must be routed to the correct department "
        "(billing, technical, shipping, account, returns, or general) and assigned a "
        "priority level (low, medium, high, urgent). This triage is typically performed "
        "manually, which introduces three persistent failure modes:"
    )
    _add_bullets(doc, [
        "Slow response times — tickets accumulate in a general queue until a human reads and re-routes them, increasing time-to-first-response.",
        "Misclassification — agents unfamiliar with edge cases route tickets to the wrong department, producing additional hand-offs and customer frustration.",
        "Inconsistent prioritisation — urgency is subjective, so one agent's high is another agent's medium, undermining service-level agreements.",
    ])
    _add_paragraph(
        doc,
        "SmartTicket addresses these problems by treating triage as a supervised "
        "learning task with two prediction targets and by enforcing explainability so "
        "that the system can be trusted, audited, and corrected by human operators."
    )

    # 3. Objectives
    _add_heading(doc, "3. Project Objectives", level=1)
    _add_paragraph(doc, "The project pursues five concrete objectives:")
    _add_bullets(doc, [
        "Construct a reproducible data pipeline that ingests, cleans, and normalises raw support tickets stored across three relational tables.",
        "Engineer informative features from heterogeneous data sources, combining textual statistics, derived customer-behaviour ratios, and one-hot encoded categorical attributes.",
        "Implement and compare two feature-selection methodologies — a greedy heuristic and a genetic algorithm — to identify a compact, high-signal feature subset.",
        "Train and evaluate ensemble classifiers (voting and stacking) that improve robustness over any individual base model.",
        "Provide post-hoc, model-agnostic explanations for every prediction using LIME and SHAP, supporting agent review and continuous quality assurance.",
    ])

    # 4. Data
    _add_heading(doc, "4. Data Sources and Schema", level=1)
    _add_paragraph(
        doc,
        "The dataset is generated synthetically but engineered to require realistic "
        "preprocessing. It resides in a SQLite database (smartticket.db) consisting of "
        "three normalised tables, joined on ticket_id."
    )

    _add_heading(doc, "4.1 Table Overview", level=2)
    _add_table(
        doc,
        headers=["Table", "Rows", "Description"],
        rows=[
            ["tickets", "~2,500",
             "Core ticket data — free-text description, channel, product category, region, plus the two target labels (department, priority)."],
            ["customer_metrics", "~2,500",
             "Customer history — account age, order count, total spend, returns, loyalty tier, satisfaction score."],
            ["ticket_metadata", "~2,500",
             "Operational metadata — response time, attachments, replies, escalation flags, sentiment score, raw word count."],
        ],
        col_widths_cm=[3.5, 2.5, 10.5],
    )

    _add_heading(doc, "4.2 Built-In Data Messiness", level=2)
    _add_paragraph(
        doc,
        "To force the pipeline to perform meaningful preprocessing, the synthetic "
        "generator deliberately injects realistic data quality problems:"
    )
    _add_bullets(doc, [
        "Approximately 5 percent of ticket texts are NULL, and approximately 3 percent of numeric columns contain NaN values.",
        "Textual noise includes mixed casing, repeated punctuation, embedded URLs and email addresses, HTML fragments, double spaces, and stray newlines.",
        "Numeric outliers exceed valid ranges, for example a total_spent value of 999,999, negative response_time_hours, and account_age_days of 99,999.",
        "Categorical columns are inconsistent in casing, for example Email, email, and EMAIL all appearing in the same column.",
        "Approximately 1 percent of records share duplicate ticket_id values.",
    ])

    _add_heading(doc, "4.3 Inputs and Outputs", level=2)
    _add_paragraph(
        doc,
        "Inputs: 24 raw columns spanning ticket text, customer metrics, and ticket "
        "metadata. Outputs: predicted department (6 classes) and predicted priority "
        "(4 levels), each accompanied by LIME and SHAP attribution scores."
    )

    # 5. Methodology
    _add_heading(doc, "5. Methodology", level=1)

    _add_heading(doc, "5.1 Preprocessing", level=2)
    _add_paragraph(
        doc,
        "Three parallel preprocessing tracks run on the joined dataset:"
    )
    _add_bullets(doc, [
        "Text — null imputation, lowercasing, removal of URLs, HTML, and emails, stopword removal, and lemmatisation, followed by TF-IDF vectorisation with 500 features.",
        "Numeric — median imputation for missing values, clipping of values outside valid ranges, and standard scaling.",
        "Categorical — casing normalisation followed by one-hot encoding of channel, product_category, region, and loyalty_tier.",
    ])

    _add_heading(doc, "5.2 Feature Engineering", level=2)
    _add_paragraph(
        doc,
        "Derived features are computed before model fitting, including word_count, "
        "char_count, avg_word_len, sentence_count, spend_per_order, return_rate, "
        "ticket_rate, and an order_recency_score. The combined feature matrix unites "
        "TF-IDF vectors, scaled numeric features, and one-hot categorical features into "
        "a single sparse representation."
    )

    _add_heading(doc, "5.3 Feature Selection", level=2)
    _add_table(
        doc,
        headers=["Method", "Approach", "Trade-off"],
        rows=[
            ["Greedy Forward (Heuristic)",
             "Iteratively adds the single feature that yields the largest gain in KNN accuracy until a budget is reached.",
             "Deterministic and fast, but may miss beneficial feature interactions."],
            ["Genetic Algorithm",
             "Evolutionary search over binary feature masks using tournament selection, single-point crossover, and bit-flip mutation.",
             "Stochastic and able to discover non-obvious feature subsets, at higher computational cost."],
        ],
        col_widths_cm=[4.5, 6.5, 5.5],
    )

    _add_heading(doc, "5.4 Classification", level=2)
    _add_paragraph(
        doc,
        "Two ensemble strategies are implemented to improve robustness over single-model "
        "baselines:"
    )
    _add_bullets(doc, [
        "Voting Ensemble — combines KNN, Decision Tree, Random Forest, Logistic Regression, and SVM via hard voting, soft voting, and accuracy-weighted soft voting.",
        "Stacking Ensemble — uses KNN, Decision Tree, Random Forest, and SVM as base estimators and Logistic Regression as a meta-learner trained on 5-fold cross-validated out-of-fold predictions.",
    ])

    _add_heading(doc, "5.5 Explainability", level=2)
    _add_paragraph(
        doc,
        "Every classification is paired with two model-agnostic explanations. LIME fits "
        "a local linear surrogate on word-level perturbations of the ticket text, while "
        "SHAP computes Shapley values over the TF-IDF feature space using a Kernel "
        "Explainer. The two methods cross-validate each other: when both agree on the "
        "top-contributing words the explanation is treated as robust; disagreement "
        "flags the prediction for closer review."
    )

    # 6. System Architecture
    _add_heading(doc, "6. System Architecture", level=1)
    _add_paragraph(
        doc,
        "The pipeline is organised as a directed flow from raw data through "
        "preprocessing, feature engineering, feature selection, ensemble modelling, "
        "and explainability, terminating in a final prediction enriched with attributions."
    )
    arch = (
        "Raw Support Tickets (SQLite, 3 tables)\n"
        "        |\n"
        "        v\n"
        "Preprocessing  -->  Text  |  Numeric  |  Categorical\n"
        "        |\n"
        "        v\n"
        "Feature Engineering (derived ratios, text statistics, TF-IDF)\n"
        "        |\n"
        "        v\n"
        "Feature Selection (Greedy Forward  +  Genetic Algorithm)\n"
        "        |\n"
        "        v\n"
        "Ensemble Modelling (Voting  +  Stacking)\n"
        "        |\n"
        "        v\n"
        "Post-Hoc Explainability (LIME  +  SHAP)\n"
        "        |\n"
        "        v\n"
        "Final Classification: department (6 classes) + priority (4 levels)"
    )
    p = doc.add_paragraph()
    run = p.add_run(arch)
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    # 7. Evaluation
    _add_heading(doc, "7. Evaluation Plan", level=1)
    _add_paragraph(
        doc,
        "Model quality is assessed against a held-out test split using accuracy as the "
        "primary metric, supplemented by per-class precision, recall, and F1-score to "
        "expose imbalance effects. Feature-selection methods are compared on both "
        "achieved accuracy and the size of the selected feature subset, since a "
        "smaller subset is preferable when accuracy is comparable. Explanation quality "
        "is assessed qualitatively by reviewing LIME and SHAP agreement on a sample of "
        "predictions."
    )
    _add_table(
        doc,
        headers=["Method", "Accuracy", "Number of Features"],
        rows=[
            ["Baseline (all features + TF-IDF)", "0.6452", "357"],
            ["Heuristic (Greedy Forward)",       "0.8116", "8"],
            ["Genetic Algorithm",                "0.7026", "21"],
        ],
        col_widths_cm=[7.0, 4.5, 5.0],
    )
    _add_paragraph(
        doc,
        "Preliminary results indicate that both feature-selection strategies outperform "
        "the full-feature baseline, with the greedy heuristic achieving the strongest "
        "accuracy using only eight features — a clear illustration that a compact, "
        "well-chosen feature set can outperform a much larger noisy one for distance-"
        "based classifiers such as KNN.",
        italic=True,
    )

    # 8. Tools and Technologies
    _add_heading(doc, "8. Tools and Technologies", level=1)
    _add_table(
        doc,
        headers=["Category", "Technology"],
        rows=[
            ["Programming Language", "Python 3.11+"],
            ["Data Storage",         "SQLite"],
            ["Data Manipulation",    "pandas, numpy, scipy"],
            ["Machine Learning",     "scikit-learn"],
            ["Natural Language",     "NLTK"],
            ["Explainability",       "LIME, SHAP"],
            ["Visualisation",        "matplotlib, seaborn, plotly"],
            ["Application Layer",    "Streamlit"],
            ["Notebook Environment", "Jupyter"],
        ],
        col_widths_cm=[5.0, 11.5],
    )

    # 9. Deliverables
    _add_heading(doc, "9. Expected Deliverables", level=1)
    _add_bullets(doc, [
        "A documented SQLite database generator that produces a realistic, intentionally noisy synthetic dataset.",
        "A reproducible pipeline available in both script and notebook form, executing end-to-end without manual intervention.",
        "Implementations of two feature-selection methods together with a comparative evaluation report.",
        "Voting and stacking ensemble classifiers, evaluated against a single-model baseline.",
        "An explainability module producing LIME and SHAP attributions for every prediction, plus aggregate global feature importances.",
        "A Streamlit dashboard that allows interactive ticket inspection and visualisation of model decisions.",
        "Final written report and presentation summarising methodology, results, and lessons learned.",
    ])

    # 10. Success Criteria
    _add_heading(doc, "10. Success Criteria", level=1)
    _add_bullets(doc, [
        "The pipeline runs end-to-end from raw database to classification results without manual intervention.",
        "Both feature-selection methods produce competitive accuracy with significantly fewer features than the full baseline.",
        "Voting and stacking ensembles outperform the single-model baseline on the held-out test set.",
        "Every classification decision is accompanied by word-level attributions from both LIME and SHAP.",
        "All preprocessing steps are logged with before-and-after statistics for auditability.",
    ])

    # 11. Conclusion
    _add_heading(doc, "11. Conclusion", level=1)
    _add_paragraph(
        doc,
        "SmartTicket demonstrates how a disciplined combination of data engineering, "
        "feature selection, ensemble learning, and explainability can convert a noisy, "
        "heterogeneous support-ticket stream into a reliable triage system. By coupling "
        "high accuracy with transparent, word-level explanations, the project produces "
        "a pipeline that is not only effective but also auditable — an essential "
        "property for systems that influence customer-facing decisions."
    )

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
